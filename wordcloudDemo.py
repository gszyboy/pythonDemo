#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# @Date    : 2019-01-08 14:39:46
# @Author  : 如风 (26982310@qq.com)
# @Link    : http://blog.youmaku.top
# @Version : $Id$
#import wordcloud
import jieba
import jieba.posseg as psg
from wordcloud import WordCloud,ImageColorGenerator
from collections import Counter
import matplotlib.pyplot as plt
from PIL import Image
import numpy as np
import random
import docx
font = r'C:\Windows\Fonts\msyh.ttc'
file = docx.Document('d:/DevCode/一些测试文件/甘肃省档案工作规范化管理办法的通知.docx')
file = docx.Document('d:/DevCode/一些测试文件/档案库房智慧感知平台解决方案模板20170908v1.0.docx')
#file = docx.Document('D:\\OneDrive\\个人项目\\张掖慈善协会\\张掖市慈善协会网站询价文件-最终.docx')
textString = ''
for word in file.paragraphs:
    textTitle = file.paragraphs[0].text
    if word.text != '':
        textString += word.text

result=jieba.cut(textString,cut_all = True)
santi_words = [x for x in result if len(x) >= 2]
#print(santi_words)
c = Counter(santi_words)
new_c = dict(c)
str_c = " ".join(santi_words)
#print("切分结果:  "+",".join(result))
a = []
b = []
c = []
d = []
print('result type:',type(result))

s =psg.cut(textString)
lists=[(v.word,v.flag) for v in s]
print('lists type:',type(lists))
for v in lists:
    # if v[1].startswith('n') and len(v[0])>=2:
    #     a.append(v[0])
    # if v[1].startswith('a'):
    #     a.append(v[0])
    # if v[1].startswith('v'):
    #     a.append(v[0])
    if len(v[0]) >=2:
        a.append(v[0])
#print(a)



c = Counter(a)
for k,v in c.items():
    d.append((k,v))
    #print(k)
    #print(v)
#print([dict(c)])
# print(b)
def grey_color_func(word, font_size, position, orientation, random_state=None,
                    **kwargs):
    return "hsl(0, 0%%, %d%%)" % random.randint(60, 100)

cloud_text=" ".join(a)
#print(cloud_text)
print(type(cloud_text))
img_array = np.array(Image.open('D:\\DevCode\\awesome-python3-webapp\\timg.jpg'))
img_color = np.array(Image.open('D:\\DevCode\\awesome-python3-webapp\\2.jpg'))
wc = WordCloud(
    background_color="white", #背景颜色
    max_words=200, #显示最大词数
    font_path=font,  #使用字体
    min_font_size=15,
    max_font_size=50, 
    # width=600,  #图幅宽度
    # height=500,
    margin=10,
    scale=5,#值越大越清晰
    prefer_horizontal=1,
    #默认值0.90，浮点数类型。表示在水平如果不合适，就旋转为垂直方向，水平放置的词数占0.9？
    mask=img_array,
    random_state=3
    )

# 根据图片生成词云颜色
image_colors = ImageColorGenerator(img_color)

#wc.generate_from_text(str_c)
wc.generate_from_frequencies(new_c)
#wc.to_file("pic.png")
default_colors = wc.to_array()
wc.recolor(color_func=image_colors, random_state=3)
plt.imshow(wc, interpolation="bilinear")
plt.figure()
plt.title("Default colors")
plt.imshow(default_colors, interpolation="bilinear")

plt.axis("off") #不显示坐标尺寸
plt.show()