#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# @Date    : 2019-01-09 09:47:44
# @Author  : 如风 (26982310@qq.com)
# @Link    : http://blog.youmaku.top
# @Version : $Id$
import os
import docx

class fetchDocx(object):
    """
    读取Docx文件
    filepath:文件路径
    filename:文件名称
    """
    def __init__(self,filepath,filename):
        super(fetchDocx, self).__init__()
        self.path = os.path.join(filepath,filename)
        self.filepath = filepath
        self.filename = filename

    def read(self):
        '''读取docx文件的内容'''
        textString = ''
        textTitle = ''
        state = 0
        msg = ''
        info = dict(filename=self.filename)
        text={}
        if self.__extension(self.filename):
            try: 
                file = docx.Document(self.path)
                for word in file.paragraphs:
                    textTitle = file.paragraphs[0].text
                    if word.text != '':
                        textString += word.text
                state = 1
                msg = '成功'            
                text = dict(textTitle=textTitle,textString=textString)
                return dict(state=state,msg=msg,info=info,text=text)
            except Exception as e:
                msg = '错误:%s' % e
                #return dict(state=state,msg=msg,info=info,text=text)
                raise e
        else:
            msg = '文件格式错误,不是docx格式!'
            return dict(state=state,msg=msg,info=info,text=text)

    def __extension(self,filename):
        name,extension=os.path.splitext(filename)
        return extension.lower().startswith('.docx')



